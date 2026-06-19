"""Apply verified per-chapter edits to a COPY of the graduation book docx.

Producer/applier split: edits_ch{1..6}.py each define EDITS (para + cell ops),
re-anchored to the real docx wording. This script applies them all to a copy,
preserving the original, and logs every edit as APPLIED / MISMATCH / MISS.
"""
import shutil
import docx

SRC = "Swarm_Robot_System_Graduation_Book_FORMATTED.docx"
DST = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"


def collapse_replace(p, find, replace):
    """Replace `find`->`replace` within a single paragraph, run-aware.

    `find` may span runs. We rebuild the paragraph text into run[0] (keeping
    run[0]'s formatting) and blank the rest. Returns occurrences replaced.
    """
    runs = p.runs
    full = "".join(r.text for r in runs)
    if find not in full:
        return 0
    n = full.count(find)
    new_full = full.replace(find, replace)
    if runs:
        runs[0].text = new_full
        for r in runs[1:]:
            r.text = ""
    else:
        p.add_run(new_full)
    return n


def set_cell(cell, new_text):
    """Set a table cell's text, keeping the first paragraph/run formatting."""
    p0 = cell.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = new_text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(new_text)
    # remove any extra paragraphs in the cell
    for extra in cell.paragraphs[1:]:
        el = extra._element
        el.getparent().remove(el)


def main():
    shutil.copyfile(SRC, DST)
    d = docx.Document(DST)

    applied = 0
    warnings = []

    for ch in range(1, 7):
        ns = {}
        exec(open("edits_ch%d.py" % ch, encoding="utf-8").read(), ns)
        edits = ns["EDITS"]
        print("\n=== Chapter %d: %d edits ===" % (ch, len(edits)))
        for e in edits:
            note = e.get("note", "")
            if e["type"] == "para":
                find, repl = e["find"], e["replace"]
                expect = e.get("expect", 1)
                hits = 0
                for p in d.paragraphs:
                    hits += collapse_replace(p, find, repl)
                if hits == 0:
                    warnings.append("MISS  ch%d: %s :: find not found: %r" % (ch, note, find[:60]))
                    print("  MISS     | %s" % note)
                elif hits != expect:
                    warnings.append("COUNT ch%d: %s :: expected %d got %d" % (ch, note, expect, hits))
                    print("  APPLIED* | (%d/%d) %s" % (hits, expect, note))
                    applied += 1
                else:
                    print("  APPLIED  | %s" % note)
                    applied += 1
            elif e["type"] == "cell":
                ti, r, c = e["table"], e["row"], e["col"]
                try:
                    cell = d.tables[ti].rows[r].cells[c]
                    set_cell(cell, e["new_text"])
                    print("  APPLIED  | T%d[%d,%d] %s" % (ti, r, c, note))
                    applied += 1
                except Exception as ex:
                    warnings.append("CELLERR ch%d T%d[%d,%d]: %s (%s)" % (ch, ti, r, c, note, ex))
                    print("  CELLERR  | T%d[%d,%d] %s" % (ti, r, c, note))

    d.save(DST)
    print("\n================ SUMMARY ================")
    print("Edits applied: %d" % applied)
    print("Warnings: %d" % len(warnings))
    for w in warnings:
        print("  " + w)
    print("Saved: %s" % DST)


if __name__ == "__main__":
    main()

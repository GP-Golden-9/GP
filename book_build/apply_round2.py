# -*- coding: utf-8 -*-
"""Round-2 edits from the author's confirmations:
  #1 Alpha HAS an Arduino Mega + encoders present-but-inactive (reverse my over-correction)
  #3 Alpha motors are JGB37-520 (12 Vdc, 107 rpm, 4.8 kg.cm) with Hall encoders
  #7 remove Beta's manipulation servo + its dedicated BEC buck (Gamma servo kept, pending)
  #8 fuses 15 A & 10 A; capacitors 1000 uF
  #11 camera = Logitech C270 (specific name)
"""
import shutil
import docx

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
shutil.copyfile(F, "Swarm_Robot_System_Graduation_Book_CORRECTED.prev2.docx")
d = docx.Document(F)


def collapse(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)


def rep(find, repl):
    n = 0
    for p in d.paragraphs:
        full = "".join(r.text for r in p.runs)
        if find in full:
            collapse(p, full.replace(find, repl))
            n += 1
    return n


def set_cell(cell, text):
    p0 = cell.paragraphs[0]
    collapse(p0, text)
    for ex in cell.paragraphs[1:]:
        ex._element.getparent().remove(ex._element)


def set_cell_by_match(table, col0_contains, col_idx, text):
    for row in table.rows:
        if col0_contains in row.cells[0].text:
            set_cell(row.cells[col_idx], text)
            return 1
    return 0


def remove_row_by_match(table, col0_contains):
    for row in list(table.rows):
        if col0_contains in row.cells[0].text:
            table._tbl.remove(row._tr)
            return 1
    return 0


EDITS = [
    # ---- #1/#3 Alpha Mega + JGB37-520 motors + encoders present-but-inactive ----
    ("P189 encoders",
     "Alpha has no wheel encoders; localization relies entirely on scan-based odometry (rf2o laser odometry feeding slam_toolbox).",
     "Alpha's four JGB37-520 geared motors carry Hall encoders, but these are currently inactive; localization relies on scan-based odometry (rf2o laser odometry feeding slam_toolbox)."),
    ("P743 drop unverified no-IMU clause",
     "successive LiDAR scans; Alpha carries no IMU.",
     "successive LiDAR scans."),
    ("P744 encoders inactive",
     "laser-based (rf2o scan-matching); Alpha has no wheel encoders.",
     "laser-based (rf2o scan-matching); Alpha's wheel encoders are present but currently inactive."),
    ("P898 motor model",
     "Alpha was equipped with four DC gear-motors (their wheel encoders are present but unused, since Alpha localizes via rf2o laser scan-matching), and Beta was equipped with four 25GA370 motors featuring integrated 408 CPR quadrature encoders.",
     "Alpha was equipped with four JGB37-520 geared motors (12 Vdc, 107 rpm, 4.8 kg.cm) whose integrated Hall encoders are present but currently inactive (Alpha localizes via rf2o laser scan-matching), and Beta was equipped with four 25GA370 motors featuring integrated 408 CPR quadrature encoders."),
    ("P1130 Alpha",
     "Alpha has no Arduino Mega and no wheel encoders; it derives its motion estimate from rf2o laser odometry (scan-matching on the RPLidar stream) rather than dead reckoning.",
     "Alpha uses an Arduino Mega 2560 for low-level motor control and carries Hall encoders on its four JGB37-520 motors that are present but currently inactive; it derives its motion estimate from rf2o laser odometry (scan-matching on the RPLidar stream) rather than wheel dead reckoning."),

    # ---- #11 camera = Logitech C270 (Beta-specific spec mentions) ----
    ("P1130 Beta list (camera + drop servo)",
     "a USB camera, a GY-87 IMU, 25GA370 DC motors with 408 CPR encoders, a water pump, and an arm servo.",
     "a Logitech C270 USB camera, a GY-87 IMU, 25GA370 DC motors with 408 CPR encoders, and a water pump."),
    ("P190 Beta list (camera + drop servo)",
     "a USB camera, a GY-87 IMU, four 25GA370 encoder motors (408 CPR), HC-SR04 ultrasonic sensors, a water pump, and a servo mechanism.",
     "a Logitech C270 USB camera, a GY-87 IMU, four 25GA370 encoder motors (408 CPR), HC-SR04 ultrasonic sensors, and a water pump."),
    ("P190 intervention tools (drop servo)",
     "carries fire intervention tools (pump and servo)",
     "carries a fire-intervention water pump"),
    ("P290 intervener (drop servo)",
     "fire intervention via pump and servo, encoder-based navigation",
     "fire intervention via water pump, encoder-based navigation"),
    ("P611 camera",
     "integrating the USB camera,",
     "integrating the Logitech C270 USB camera,"),
    ("P649 camera",
     "from the onboard USB camera",
     "from the onboard Logitech C270 USB camera"),
    ("P703 camera",
     "Sensor: USB camera configured for 640x480",
     "Sensor: Logitech C270 USB camera configured for 640x480"),
    ("P751 camera",
     "Captures images from the USB camera",
     "Captures images from the Logitech C270 USB camera"),
    ("P990 camera",
     "a generic USB camera (device index 0,",
     "a Logitech C270 USB camera (device index 0,"),

    # ---- #7 remove Beta servo + its dedicated BEC ----
    ("P928 remove BEC/servo sentence",
     "The second converter (a dedicated BEC buck) regulates the voltage to 6.0 V / 3 A for the servo rail (Beta's arm servo and Gamma's sensor-mount servo); note that Beta is skid-steer and has no steering servo. ",
     ""),
    ("P1223 remove BEC/servo clause",
     "; and a separate BEC rail for the servo) mitigated motor-induced voltage dips.",
     ") mitigated motor-induced voltage dips."),

    # ---- #8 fuses & capacitors ----
    ("P929 caps",
     "Capacitive filtering was applied at the output stage of both buck converters to suppress",
     "Capacitive filtering using 1000 µF capacitors was applied at the buck-converter output to suppress"),
    ("P929 fuses",
     "inline fuses were installed on the main power distribution line",
     "inline fuses rated at 15 A and 10 A were installed on the main power distribution line"),
]

print("== paragraph edits ==")
warn = []
for note, find, repl in EDITS:
    n = rep(find, repl)
    print("  %-42s applied=%d" % (note, n))
    if n != 1:
        warn.append("%s -> %d" % (note, n))

print("== table edits ==")
t8, t9, t11 = d.tables[8], d.tables[9], d.tables[11]
print("  T8 Mega function   :", set_cell_by_match(t8, "Real-Time Control Unit", 3,
      "Motor control and low-level I/O (encoder inputs wired but currently inactive)"))
print("  T8 encoders spec   :", set_cell_by_match(t8, "Wheel Encoders", 2,
      "Hall encoders integrated in Alpha's four JGB37-520 geared motors (12 Vdc, 107 rpm, 4.8 kg.cm), currently inactive; Beta uses 25GA370, 408 CPR quadrature"))
print("  T8 encoders func   :", set_cell_by_match(t8, "Wheel Encoders", 3,
      "Alpha: present but inactive (localization by scan-matching). Beta: wheel odometry"))
print("  T9 encoders spec   :", set_cell_by_match(t9, "Wheel Encoders", 1,
      "JGB37-520 Hall encoders (Alpha, present but inactive); 25GA370, 408 CPR quadrature (Beta)"))
print("  T11 remove BEC row :", remove_row_by_match(t11, "Buck Converter 2"))
print("  T11 fuse spec      :", set_cell_by_match(t11, "Inline Fuse", 1, "15 A and 10 A inline fuses"))
print("  T11 cap spec       :", set_cell_by_match(t11, "Capacitive Filters", 1, "1000 µF output capacitors"))

d.save(F)
print("Saved:", F)
if warn:
    print("WARNINGS:", warn)

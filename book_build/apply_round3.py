# -*- coding: utf-8 -*-
"""Round-3 edits from author's confirmations:
  #1 Alpha has NO IMU -> leave as-is (no change)
  #2 Gamma servo cancelled -> remove (P191, P1034, P1130); Gamma keeps ONE ultrasonic
  #3 Per-robot batteries: Alpha 5200 mAh, Beta 2200 mAh, Gamma 3x 18650 Li-ion
  #4 All regulators = XL4015 5 A adjustable (SKU DC102) -> reconcile tables + prose
  #5 Figures 5.1/5.2 keep placeholders (already inserted) -> no change
"""
import shutil
import docx

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
shutil.copyfile(F, "Swarm_Robot_System_Graduation_Book_CORRECTED.prev3.docx")
d = docx.Document(F)


def collapse(p, new):
    if p.runs:
        p.runs[0].text = new
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new)


def rep(find, repl):
    n = 0
    for p in d.paragraphs:
        full = "".join(r.text for r in p.runs)
        if find in full:
            collapse(p, full.replace(find, repl))
            n += 1
    return n


def set_cell(cell, text):
    p0 = cell.paragraphs[0]
    collapse(p0, text)
    for ex in cell.paragraphs[1:]:
        ex._element.getparent().remove(ex._element)


EDITS = [
    # ---- #2 Gamma servo removal ----
    ("P191 Gamma servo",
     "an MPU6050 IMU, a buzzer, and a servo mechanism.",
     "an MPU6050 IMU, and a buzzer."),
    ("P1034 Gamma /servo endpoint",
     ", and /servo?deg=S for steering the sensor mount.",
     "."),
    ("P1130 Gamma tail",
     "Gamma was an ESP32-based platform equipped with an MQ-5 gas sensor, HC-SR04 ultrasonic sensor, buzzer, and servo.",
     "Gamma was an ESP32-based platform equipped with an MQ-5 gas sensor, a single HC-SR04 ultrasonic sensor, and a buzzer."),

    # ---- #3 per-robot batteries (prose) ----
    ("P927 per-robot battery",
     "The system utilizes a 3S lithium-polymer (LiPo) battery with a nominal voltage of 11.1 V and a capacity of 5000 mAh, providing approximately 55 Wh of usable energy. This configuration offers an optimal balance between energy density, weight, and discharge capability, making it suitable for prolonged mobile operation without significantly increasing the robot payload.",
     "Battery capacity is sized per robot. Alpha (the mapping platform) uses a 3S LiPo pack (11.1 V, 5200 mAh, ~57.7 Wh); Beta (the intervention platform) uses a smaller 3S LiPo pack (11.1 V, 2200 mAh, ~24.4 Wh); and Gamma (the ESP32 inspector) runs on three series lithium-ion 18650 cells (~11.1 V nominal). This per-platform sizing balances energy density, weight, and discharge capability against each robot's payload and runtime needs."),

    # ---- #4 XL4015 regulator (prose) ----
    ("P928 buck XL4015",
     "Logic and auxiliary systems are isolated using independent step-down (buck) converters. The first converter targets ~5.25 V at the Pi pins (set higher to offset the ~0.4-0.5 V delivery-path drop) at 5 A, dedicated to powering the Raspberry Pi and Arduino Mega logic.",
     "Logic is isolated using an adjustable step-down (buck) converter. The system standardizes on a single regulator model across all three robots - the XL4015 5 A adjustable buck (4.8-38 V input, 1.25-36 V output, SKU DC102) - set on each mobile robot to target ~5.25 V at the Pi pins (raised to offset the ~0.4-0.5 V delivery-path drop) to power the Raspberry Pi and Arduino Mega logic."),
    ("P1277 dual->single buck",
     "The platform uses a 3S LiPo battery with dual independent buck converters intended to isolate sensitive electronics, such as the Raspberry Pi, from motor-induced voltage brownouts.",
     "The platform uses a 3S LiPo battery with an adjustable XL4015 buck converter intended to isolate sensitive electronics, such as the Raspberry Pi, from motor-induced voltage brownouts."),
]

print("== paragraph edits ==")
warn = []
for note, find, repl in EDITS:
    n = rep(find, repl)
    print("  %-30s applied=%d" % (note, n))
    if n != 1:
        warn.append("%s -> %d" % (note, n))

print("== table edits ==")
t8, t11 = d.tables[8], d.tables[11]
set_cell(t8.rows[6].cells[2], "11.1 V, 5200 mAh, ~57.7 Wh")
print("  T8 battery (Alpha) -> set")
set_cell(t8.rows[7].cells[2], "XL4015 5 A adjustable buck, set ~5.25 V for Pi/Mega logic; motors battery-direct via L298N")
print("  T8 power-regulation -> set")
set_cell(t11.rows[1].cells[1], "3S LiPo, 11.1 V nominal; per robot: Alpha 5200 mAh (~57.7 Wh), Beta 2200 mAh (~24.4 Wh)")
print("  T11 main battery -> set")
set_cell(t11.rows[2].cells[1], "XL4015 5 A adjustable buck (4.8-38 V in, 1.25-36 V out, SKU DC102): 11.1 V down to ~5.25 V at the Pi pins (raised to offset the ~0.4-0.5 V path drop)")
print("  T11 buck1 -> set")

d.save(F)
print("Saved:", F)
if warn:
    print("WARNINGS:", warn)

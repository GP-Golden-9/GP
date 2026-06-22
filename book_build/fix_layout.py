# -*- coding: utf-8 -*-
"""Fix layout problems in the friend-edited book:
  1. Restore Figure 2.1 (montage of the 5 hardware-component photos the friend deleted).
  2. Remove the duplicated GitHub hyperlink in the 4.2 Software Implementation paragraph.
  3. Collapse runs of stray empty paragraphs (the big blank gaps) to a single blank,
     preserving any paragraph that carries a page break or section break.
Page numbers are recomputed separately afterwards.
"""
import os, docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from PIL import Image

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"

# ---------- 1. build the Figure 2.1 montage ----------
srcs = ["docs/book_images/chapter2/figure_2.1_main_hardware_components_%s.png" % c for c in "abcde"]
cell_w, cell_h, gap, pad = 380, 300, 24, 24
cols, rows = 3, 2
W = pad*2 + cols*cell_w + (cols-1)*gap
H = pad*2 + rows*cell_h + (rows-1)*gap
montage = Image.new("RGB", (W, H), (255, 255, 255))
for idx, s in enumerate(srcs):
    im = Image.open(s).convert("RGB")
    im.thumbnail((cell_w, cell_h))
    r, c = divmod(idx, cols)
    # center the 5th image under the gap of row 2 (cols=3, row2 has 2 -> center them)
    if r == 1:
        # shift the two bottom images to be centered
        x0 = pad + (cell_w + gap)//2 + c*(cell_w+gap)
    else:
        x0 = pad + c*(cell_w+gap)
    y0 = pad + r*(cell_h+gap)
    x = x0 + (cell_w - im.width)//2
    y = y0 + (cell_h - im.height)//2
    montage.paste(im, (x, y))
MON = "figures_png/_fig21_montage.png"
montage.save(MON)
print("montage built:", MON, montage.size)

d = docx.Document(F)

def has_img(p): return p._p.find('.//' + qn('a:blip')) is not None
def has_break(p):
    if p.paragraph_format.page_break_before: return True
    if p._p.find('.//' + qn('w:br')) is not None:
        for br in p._p.iter(qn('w:br')):
            if br.get(qn('w:type')) == 'page': return True
    if p._p.find('.//' + qn('w:sectPr')) is not None: return True
    return False

bs = next(i for i, p in enumerate(d.paragraphs) if p.text.strip() == "Chapter 1: Introduction" and i > 120)

# ---------- restore Figure 2.1 image (insert before its caption) ----------
cap = next(p for i, p in enumerate(d.paragraphs) if i > bs and p.text.strip().startswith("Figure 2.1") and "Main hardware" in p.text)
img_p_el = OxmlElement('w:p'); cap._p.addprevious(img_p_el)
img_p = Paragraph(img_p_el, cap._parent)
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.paragraph_format.keep_with_next = True
img_p.paragraph_format.space_before = Pt(6); img_p.paragraph_format.space_after = Pt(2)
img_p.add_run().add_picture(MON, width=Inches(5.6))
print("Figure 2.1 montage inserted before caption")

# ---------- 2. remove duplicate GitHub hyperlink ----------
fixed_url = 0
for p in d.paragraphs:
    hls = p._p.findall(qn('w:hyperlink'))
    for h in hls:
        if "github.com/GP-Golden" in "".join(t.text or "" for t in h.iter(qn('w:t'))):
            p._p.remove(h); fixed_url += 1
# clean any resulting double period / stray ").." in those paragraphs
for p in d.paragraphs:
    if "GP)." in p.text and p.runs:
        full = "".join(r.text for r in p.runs)
        new = full.replace("GP).https://github.com/GP-Golden-9/GP).", "GP).").replace("GP)..", "GP).")
        if new != full:
            p.runs[0].text = new
            for r in p.runs[1:]: r.text = ""
print("duplicate hyperlinks removed:", fixed_url)

# ---------- 3. collapse stray empty-paragraph runs (keep 1, preserve breaks) ----------
removed = 0
run = []
def flush(run):
    global removed
    # keep the first; if any in run has a break, keep that one instead; delete the others
    if len(run) <= 1: return
    keep = run[0]
    for p in run:
        if has_break(p): keep = p; break
    for p in run:
        if p is keep: continue
        if has_break(p): continue          # never delete a break-bearing paragraph
        p._p.getparent().remove(p._p); removed += 1

paras = list(d.paragraphs)
i = bs
cur = []
for p in paras[bs:]:
    empty = (not p.text.strip()) and not has_img(p)
    if empty:
        cur.append(p)
    else:
        flush(cur); cur = []
flush(cur)
print("empty paragraphs removed:", removed)

d.save(F)
print("saved:", F)

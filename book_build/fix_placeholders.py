# -*- coding: utf-8 -*-
"""Regenerate the two Chapter-5 placeholder images with their CORRECT (renumbered)
figure numbers baked in, and swap them into the docx. (They were generated as 5.1/5.2
before figure renumbering; captions are now 5.2 and 5.6.)"""
import os, docx
from docx.shared import Inches
from PIL import Image, ImageDraw, ImageFont

def make_placeholder(path, fig_no, caption):
    W, H = 1280, 720
    img = Image.new("RGB", (W, H), (242, 244, 247)); dr = ImageDraw.Draw(img)
    dr.rectangle([8, 8, W - 8, H - 8], outline=(150, 158, 170), width=4)
    for x in range(40, W - 40, 28):
        dr.line([(x, 40), (x + 14, 40)], fill=(190, 196, 205), width=3)
        dr.line([(x, H - 40), (x + 14, H - 40)], fill=(190, 196, 205), width=3)
    for y in range(40, H - 40, 28):
        dr.line([(40, y), (40, y + 14)], fill=(190, 196, 205), width=3)
        dr.line([(W - 40, y), (W - 40, y + 14)], fill=(190, 196, 205), width=3)
    cx, cy = W // 2, H // 2 - 70
    dr.rectangle([cx - 110, cy - 80, cx + 110, cy + 80], outline=(120, 130, 145), width=5)
    dr.ellipse([cx - 70, cy - 50, cx - 30, cy - 10], outline=(120, 130, 145), width=5)
    dr.polygon([(cx - 90, cy + 60), (cx - 20, cy - 5), (cx + 30, cy + 35),
                (cx + 70, cy - 10), (cx + 90, cy + 60)], outline=(120, 130, 145))
    def font(sz):
        for nm in ("arialbd.ttf", "DejaVuSans-Bold.ttf"):
            try: return ImageFont.truetype(nm, sz)
            except Exception: pass
        return ImageFont.load_default()
    def center(text, y, sz, col):
        f = font(sz); w = dr.textlength(text, font=f); dr.text(((W - w) / 2, y), text, fill=col, font=f)
    center("PLACEHOLDER IMAGE", cy + 120, 46, (180, 70, 70))
    center("Figure %s" % fig_no, cy + 185, 40, (60, 70, 90))
    words, line, lines = caption.split(), "", []
    for w in words:
        if dr.textlength((line + " " + w).strip(), font=font(30)) > W - 160:
            lines.append(line); line = w
        else: line = (line + " " + w).strip()
    lines.append(line); yy = cy + 235
    for ln in lines:
        center(ln, yy, 30, (90, 98, 112)); yy += 40
    center("Replace with the actual figure when available", yy + 14, 24, (140, 148, 160))
    img.save(path)

PLACE = [
    ("5.2", "SLAM Occupancy Grid Output vs. Physical Test Environment", "figures_png/_ph_5_2.png"),
    ("5.6", "Dashboard Gas Concentration Heatmap Visualization", "figures_png/_ph_5_6.png"),
]
for num, cap, path in PLACE:
    make_placeholder(path, num, cap)

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
d = docx.Document(F)
from docx.oxml.ns import qn
def has_img(p): return p._p.find('.//' + qn('a:blip')) is not None
swapped = []
for num, cap, path in PLACE:
    capkey = "Figure %s " % num
    cap_p = next(p for p in d.paragraphs if p.text.strip().startswith(capkey) and cap[:15] in p.text)
    # image is in the adjacent paragraph (after, else before)
    nxt = cap_p._p.getnext(); prv = cap_p._p.getprevious()
    from docx.text.paragraph import Paragraph
    target = None
    for el in (nxt, prv):
        if el is not None and el.tag == qn('w:p'):
            pp = Paragraph(el, cap_p._parent)
            if has_img(pp): target = pp; break
    if target is None:
        swapped.append((num, "NO IMAGE PARA FOUND")); continue
    for r in list(target.runs): r._element.getparent().remove(r._element)
    target.add_run().add_picture(path, width=Inches(5.8))
    swapped.append((num, "ok"))
d.save(F)
print("placeholder swaps:", swapped)

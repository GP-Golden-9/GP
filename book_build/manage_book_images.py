"""Extract, review, and organize the graduation-book images; generate
placeholders for missing figures and insert them into the docx.

Outputs:
  docs/book_images/<section>/<descriptive-name>   organized real images
  docs/book_images/_placeholders/                  generated placeholder PNGs
  docs/book_images/_unused_or_duplicate/           dup/header images kept for ref
  docs/book_images/IMAGE_AUDIT.md                  present/missing audit
and inserts placeholder pictures into the two "[Insert Figure Here]" slots.
"""
import os
import zipfile
import hashlib
import shutil
import docx
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont

DOCX = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
ROOT = "docs/book_images"

# media filename -> (subfolder, descriptive name).  Primary (non-duplicate) only.
FIGURE_MAP = {
    "image1.png":  ("cover", "cover_logo_1.png"),
    "image2.png":  ("cover", "cover_logo_2.png"),
    "image22.png": ("cover", "page_header_logo.png"),
    "image3.png":  ("chapter2", "figure_2.1_main_hardware_components_a.png"),
    "image4.png":  ("chapter2", "figure_2.1_main_hardware_components_b.png"),
    "image5.png":  ("chapter2", "figure_2.1_main_hardware_components_c.png"),
    "image6.png":  ("chapter2", "figure_2.1_main_hardware_components_d.png"),
    "image7.png":  ("chapter2", "figure_2.1_main_hardware_components_e.png"),
    "image11.png": ("chapter3", "figure_3.1_multi_robot_system_configuration.png"),
    "image13.png": ("chapter3", "figure_3.2_system_architecture_comm_framework.png"),
    "image15.png": ("chapter3", "figure_3.3_cad_model_robot_chassis.png"),
    "image17.png": ("chapter3", "figure_3.4_final_assembled_mapping_explorer.png"),
    "image18.jpeg":("chapter3", "figure_3.5_software_stack_ros_node_architecture.jpeg"),
    "image19.png": ("chapter3", "figure_3.6_pyside6_desktop_dashboard.png"),
    "image21.png": ("chapter3", "figure_3.7_realtime_dashboard_interface.png"),
}

# Required figures per the List of Figures + in-text captions.
REQUIRED = [
    ("2.1", "Main hardware components", "image3.png (montage a-e)"),
    ("3.1", "Multi-robot system configuration", "image11.png"),
    ("3.2", "Multi-Robot System Architecture and Communication Framework", "image13.png"),
    ("3.3", "CAD model of the robot chassis", "image15.png"),
    ("3.4", "Final Assembled Design of the Mapping Explorer Robot", "image17.png"),
    ("3.5", "Software Stack and ROS Node Architecture", "image18.jpeg"),
    ("3.6", "Native PySide6 desktop monitoring dashboard", "image19.png"),
    ("3.7", "Real-Time Dashboard Interface", "image21.png"),
    ("5.1", "SLAM Occupancy Grid Output vs. Physical Test Environment", None),
    ("5.2", "Dashboard Gas Concentration Heatmap Visualization", None),
]

MISSING = {"5.1": "SLAM Occupancy Grid Output vs. Physical Test Environment",
           "5.2": "Dashboard Gas Concentration Heatmap Visualization"}


def make_placeholder(path, fig_no, caption):
    W, H = 1280, 720
    img = Image.new("RGB", (W, H), (242, 244, 247))
    dr = ImageDraw.Draw(img)
    dr.rectangle([8, 8, W - 8, H - 8], outline=(150, 158, 170), width=4)
    # dashed inner frame
    for x in range(40, W - 40, 28):
        dr.line([(x, 40), (x + 14, 40)], fill=(190, 196, 205), width=3)
        dr.line([(x, H - 40), (x + 14, H - 40)], fill=(190, 196, 205), width=3)
    for y in range(40, H - 40, 28):
        dr.line([(40, y), (40, y + 14)], fill=(190, 196, 205), width=3)
        dr.line([(W - 40, y), (W - 40, y + 14)], fill=(190, 196, 205), width=3)
    # simple picture icon
    cx, cy = W // 2, H // 2 - 70
    dr.rectangle([cx - 110, cy - 80, cx + 110, cy + 80], outline=(120, 130, 145), width=5)
    dr.ellipse([cx - 70, cy - 50, cx - 30, cy - 10], outline=(120, 130, 145), width=5)
    dr.polygon([(cx - 90, cy + 60), (cx - 20, cy - 5), (cx + 30, cy + 35),
                (cx + 70, cy - 10), (cx + 90, cy + 60)], outline=(120, 130, 145))

    def font(sz):
        try:
            return ImageFont.truetype("arialbd.ttf", sz)
        except Exception:
            try:
                return ImageFont.truetype("DejaVuSans-Bold.ttf", sz)
            except Exception:
                return ImageFont.load_default()

    def center(text, y, sz, col):
        f = font(sz)
        w = dr.textlength(text, font=f)
        dr.text(((W - w) / 2, y), text, fill=col, font=f)

    center("PLACEHOLDER IMAGE", cy + 120, 46, (180, 70, 70))
    center("Figure %s" % fig_no, cy + 185, 40, (60, 70, 90))
    # wrap caption
    words, line, lines = caption.split(), "", []
    for w in words:
        if dr.textlength((line + " " + w).strip(), font=font(30)) > W - 160:
            lines.append(line); line = w
        else:
            line = (line + " " + w).strip()
    lines.append(line)
    yy = cy + 235
    for ln in lines:
        center(ln, yy, 30, (90, 98, 112)); yy += 40
    center("Replace with the actual figure when available", yy + 14, 24, (140, 148, 160))
    img.save(path)


def main():
    # fresh tree
    if os.path.isdir(ROOT):
        shutil.rmtree(ROOT)
    for sub in ("cover", "chapter2", "chapter3", "chapter5", "_placeholders", "_unused_or_duplicate"):
        os.makedirs(os.path.join(ROOT, sub), exist_ok=True)

    z = zipfile.ZipFile(DOCX)
    media = {n.split("/")[-1]: z.read(n) for n in z.namelist() if n.startswith("word/media/")}
    md5 = {name: hashlib.md5(data).hexdigest() for name, data in media.items()}

    # primary md5s (those we map by name)
    primary_md5 = {md5[n] for n in FIGURE_MAP if n in md5}

    organized, duplicates = [], []
    for name, (sub, newname) in FIGURE_MAP.items():
        if name in media:
            with open(os.path.join(ROOT, sub, newname), "wb") as f:
                f.write(media[name])
            organized.append((name, sub + "/" + newname))

    # everything else -> duplicate/unused
    for name, data in sorted(media.items()):
        if name in FIGURE_MAP:
            continue
        tag = "dup-of-primary" if md5[name] in primary_md5 else "unreferenced"
        with open(os.path.join(ROOT, "_unused_or_duplicate", name), "wb") as f:
            f.write(data)
        duplicates.append((name, tag))

    # placeholders for missing figures
    placeholders = {}
    for fig, cap in MISSING.items():
        pth = os.path.join(ROOT, "_placeholders", "figure_%s_PLACEHOLDER.png" % fig)
        make_placeholder(pth, fig, cap)
        placeholders[fig] = pth
        # also drop a copy in chapter5 with descriptive name
        nm = "figure_%s_%s_PLACEHOLDER.png" % (
            fig, cap.lower().replace(" ", "_").replace(".", "").replace(",", "")[:40])
        shutil.copyfile(pth, os.path.join(ROOT, "chapter5", nm))

    # generic default template
    make_placeholder(os.path.join(ROOT, "_placeholders", "_default_placeholder.png"),
                     "X.Y", "Figure caption goes here")

    # ---- insert placeholders into the docx at the two [Insert Figure Here] slots ----
    d = docx.Document(DOCX)
    order = []  # figure number in caption order
    inserted = []
    paras = d.paragraphs
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("Figure 5.1"):
            order.append(("5.1", i))
        elif t.startswith("Figure 5.2"):
            order.append(("5.2", i))
    # the [Insert Figure Here] immediately after each caption
    for fig, cap_i in order:
        for j in range(cap_i + 1, min(cap_i + 4, len(paras))):
            if paras[j].text.strip() == "[Insert Figure Here]":
                p = paras[j]
                for r in p.runs:
                    r.text = ""
                run = (p.runs[0] if p.runs else p.add_run())
                run.add_picture(placeholders[fig], width=Inches(5.8))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted.append(fig)
                break
    d.save(DOCX)

    # ---- audit report ----
    lines = ["# Book Image Audit", "",
             "Source: `%s`  |  Embedded media files: %d (incl. duplicates)" % (DOCX, len(media)),
             "", "## Required figures (List of Figures + in-text captions)", "",
             "| Figure | Caption | Status | Source image |",
             "|--------|---------|--------|--------------|"]
    for fig, cap, src in REQUIRED:
        if fig in MISSING:
            status = "MISSING -> placeholder inserted"
            src = "_placeholders/figure_%s_PLACEHOLDER.png" % fig
        else:
            status = "present"
        lines.append("| %s | %s | %s | %s |" % (fig, cap, status, src))
    lines += ["", "## Organized real images", ""]
    for name, dest in organized:
        lines.append("- `%s` -> `%s/%s`" % (name, ROOT, dest))
    lines += ["", "## Duplicate / unreferenced media (kept for reference)", ""]
    for name, tag in duplicates:
        lines.append("- `%s` (%s)" % (name, tag))
    lines += ["", "## Folder layout", "",
              "```", "docs/book_images/",
              "  cover/            university / project / header logos",
              "  chapter2/         Figure 2.1 hardware montage (a-e)",
              "  chapter3/         Figures 3.1 - 3.7",
              "  chapter5/         Figures 5.1, 5.2 (placeholders for now)",
              "  _placeholders/    generated placeholder PNGs + default template",
              "  _unused_or_duplicate/  Word's duplicate copies + page-header logo",
              "```", "",
              "## Action needed from you", "",
              "Provide the two real Chapter-5 figures and drop them into `chapter5/`:",
              "- **Figure 5.1** - SLAM occupancy-grid output vs. the physical test area "
              "(screenshot the dashboard map next to a photo of the arena).",
              "- **Figure 5.2** - dashboard gas-concentration view when the MQ-5 alarm trips.",
              "Then tell me and I'll swap the placeholders for the real images."]
    with open(os.path.join(ROOT, "IMAGE_AUDIT.md"), "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    print("organized:", len(organized), "| duplicates/unused:", len(duplicates),
          "| placeholders inserted for:", inserted)
    print("audit written:", os.path.join(ROOT, "IMAGE_AUDIT.md"))


if __name__ == "__main__":
    main()

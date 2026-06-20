# -*- coding: utf-8 -*-
"""Stage 1 of figure integration:
- replace the 3 weak figures (old 3.2/3.5/3.6) with new generated diagrams,
- insert the other 27 new figures at their section anchors,
- renumber EVERY figure per chapter by position,
- remap in-text 'Figure X.Y' references so prose stays correct.
Outputs ordered figure list + old->new map to _figs.json for Stage 2 (List of Figures).
"""
import re, json, shutil, docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
PNG = "figures_png/%s.png"
shutil.copyfile(F, "Swarm_Robot_System_Graduation_Book_CORRECTED.prefigs.docx")
d = docx.Document(F)
BLACK = RGBColor(0, 0, 0)
DASH = "–"      # en-dash separator
TOK = "§§"  # temp number token for new figures

def body_start():
    return next(i for i, p in enumerate(d.paragraphs)
                if p.text.strip() == "Chapter 1: Introduction" and i > 120)

def find_para(key, startswith=False):
    bs = body_start()
    # tolerate headings that lost their leading number (e.g. "Problem Statement"
    # instead of "1.3 Problem Statement") by also trying the title-only form.
    title = re.sub(r"^\d+(\.\d+)*\s+", "", key)
    cands = [key] if title == key else [key, title]
    for cand in cands:
        cnorm = " ".join(cand.split())
        for i in range(bs, len(d.paragraphs)):
            t = d.paragraphs[i].text.strip()
            if startswith:
                if t.startswith(cand):
                    return d.paragraphs[i]
            elif t == cand or " ".join(t.split()) == cnorm:
                return d.paragraphs[i]
    raise KeyError(key)

def new_after(par):
    e = OxmlElement('w:p'); par._p.addnext(e)
    return Paragraph(e, par._parent)

def caption_para(p, text):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.font.name = "Cambria"; r.font.size = Pt(10)
    r.font.bold = True; r.font.color.rgb = BLACK

def image_para(p, png, width):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(png, width=Inches(width))

# (key, width, title)
REPLACE = [  # old caption "Figure N" -> new image, keep number (renumber fixes it)
    ("Figure 3.2", "b1_system_architecture", 6.3, "Overall multi-robot system architecture"),
    ("Figure 3.5", "b4_software_stack", 6.1, "Software stack and ROS 2 node architecture"),
    ("Figure 3.6", "b7_dashboard_layout", 6.3, "Operations-center dashboard layout"),
]
# anchor_spec -> list of (key, width, title) inserted in order after the anchor
INSERT = [
    ("1.1 Project Overview", False, [("e1_robot_roles", 6.2, "Robot roles and responsibilities across the fleet")]),
    ("1.3 Problem Statement", False, [("a1_problem_solution", 6.2, "Problems addressed by the Swarm Robot System")]),
    ("3.1.2 Project Timeline and Gantt Chart", False, [("e3_gantt", 6.2, "Project timeline (Gantt chart)")]),
    ("3.3.2 Functional Requirements", False, [("e5_use_case", 5.6, "System use-case diagram")]),
    ("Figure 3.2", True, [("b3_ros_island", 5.6, "ROS-island isolation and the gateway boundary"),
                          ("e2_frame_alignment", 5.0, "Shared-map frame alignment across robots")]),
    ("Figure 3.4", True, [("b5a_hardware_alpha", 5.6, "Hardware block diagram — Alpha (mapping robot)"),
                          ("b5b_hardware_beta", 5.6, "Hardware block diagram — Beta (intervention robot)"),
                          ("b5c_hardware_gamma", 5.6, "Hardware block diagram — Gamma (inspection robot)")]),
    ("Figure 3.5", True, [("e4_operating_modes", 5.2, "Robot operating-mode state diagram")]),
    ("Figure 3.6", True, [("b2_protocol_stack", 6.0, "Communication protocol channels and ports"),
                          ("c7_command_sequence", 5.6, "Command acknowledgement, de-duplication and deadman sequence")]),
    ("4.1 Hardware Implementation", False, [("b6_power_architecture", 5.7, "Power architecture and distribution")]),
    ("4.2.1 SLAM Implementation", False, [("c1_slam_pipeline", 6.0, "SLAM and mapping pipeline")]),
    ("4.2.2 Visual Hazard Detection Implementation", False, [("c2_fire_detection", 5.6, "Fire-detection pipeline")]),
    ("4.2.3 Sensor Integration", False, [("c5_odometry_fusion", 5.7, "Sensor-fusion and odometry data flow"),
                                         ("c8_stall_state", 5.3, "Stall-detection and anti-lockup state machine")]),
    ("4.3 System Integration", False, [("c3_gotofire_mission", 4.5, "Autonomous fire-response (GO-TO-FIRE) mission flow"),
                                       ("c4_scan_coverage", 5.6, "Area-coverage (SCAN-AREA) mission flow"),
                                       ("c6_deadman_safety", 5.7, "Four-layer deadman safety chain"),
                                       ("c9_boot_supervision", 5.7, "Boot and process-supervision sequence")]),
    ("5.1.2 Physical Test Environment", False, [("a2_scenario_map", 6.0, "Indoor test deployment scenario")]),
    ("5.3.2 Visual Perception Results", False, [("d4_detection_thresholds", 5.6, "Configured detection confidence thresholds")]),
    ("5.3.3 Sensor Fusion and Localization Results", False, [("d2_heading_drift", 4.7, "Heading drift before and after gyro calibration")]),
    ("5.3.4 Environmental Sensing Results", False, [("d3_gas_threshold", 5.7, "Gas-alarm threshold and hysteresis behaviour")]),
    ("5.3.5 Network and Communication Results", False, [("d1_network_kpis", 6.2, "Network KPIs versus acceptance gates"),
                                                        ("d5_pipeline_summary", 6.0, "Communication pipeline throughput and latency summary")]),
]

# 1) resolve ALL anchors to paragraph objects FIRST (indices stable as objects)
anchors = [(find_para(spec, sw), items) for (spec, sw, items) in INSERT]
repl_caps = [(find_para(old, True), key, w, t) for (old, key, w, t) in REPLACE]

# 2) replace the 3 weak figures (image in the paragraph before the caption)
for cap, key, w, title in repl_caps:
    prev = cap._p.getprevious()
    imgp = Paragraph(prev, cap._parent)
    image_para(imgp, PNG % key, w)
    oldnum = re.match(r"Figure\s+(\S+)", cap.text).group(1)
    caption_para(cap, "Figure %s %s %s" % (oldnum, DASH, title))

# 3) insert the new figures (chained after each anchor)
for anchor, items in anchors:
    cur = anchor
    for key, w, title in items:
        ip = new_after(cur); image_para(ip, PNG % key, w)
        cp = new_after(ip); caption_para(cp, "Figure %s %s %s" % (TOK, DASH, title))
        cur = cp

# 4) renumber every figure per chapter (body only)
bs = body_start()
fig_re = re.compile(r"^Figure\s+(\S+)\s+%s\s+(.+)$" % DASH)
chap = 0; cnt = {}; oldnew = {}; ordered = []
for p in d.paragraphs[bs:]:
    t = p.text.strip()
    if len(t) < 70 and (re.match(r"^Chapter\s+(\d+)", t) or re.match(r"^\d+(\.\d+)*\s+[A-Z]", t)):
        chap = int(re.match(r"^(?:Chapter\s+)?(\d+)", t).group(1)); continue
    m = fig_re.match(t)
    if m and len(t) < 150:
        cnt[chap] = cnt.get(chap, 0) + 1
        new = "%d.%d" % (chap, cnt[chap]); title = m.group(2)
        if m.group(1) != TOK:
            oldnew[m.group(1)] = new
        caption_para(p, "Figure %s %s %s" % (new, DASH, title))
        ordered.append([new, title])

# 5) remap in-text 'Figure X.Y' references in body prose (token pass = no cascade)
items = sorted(oldnew.items(), key=lambda kv: -len(kv[0]))
for p in d.paragraphs[bs:]:
    full = "".join(r.text for r in p.runs)
    if "Figure" not in full or fig_re.match(full.strip()):
        continue
    new = full
    for i, (old, _) in enumerate(items):
        new = re.sub(r"Figure\s+" + re.escape(old) + r"\b", "\x00%d\x00" % i, new)
    for i, (old, nw) in enumerate(items):
        new = new.replace("\x00%d\x00" % i, "Figure " + nw)
    if new != full:
        p.runs[0].text = new
        for r in p.runs[1:]: r.text = ""

d.save(F)
json.dump({"ordered": ordered, "oldnew": oldnew}, open("book_build/_figs.json", "w"))
print("replaced:", len(repl_caps), "| inserted:", sum(len(i) for _, i in anchors),
      "| total figures now:", len(ordered))
print("old->new number map:", oldnew)

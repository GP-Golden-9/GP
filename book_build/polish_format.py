# -*- coding: utf-8 -*-
"""Professional formatting polish for the graduation book.
- Body: Cambria 12, fully justified, 1.3 line spacing, 8pt after, widow control.
- Headings: Calibri bold, navy (#1F3864), sized by level, keep-with-next.
- Tables: navy header row (white bold, repeats), zebra data rows, padded cells,
  rows can't split across pages; captions kept with their table.
- Margins widened for comfortable reading. Title/front matter: font unified only.
Text content is NOT altered - formatting only.
"""
import re, shutil, docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
shutil.copyfile(F, "Swarm_Robot_System_Graduation_Book_CORRECTED.prefmt.docx")
d = docx.Document(F)

BODY = "Cambria"
HEAD = "Calibri"
NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x44, 0x54, 0x6A)
HDR_FILL = "1F3864"
ZEBRA = "EEF2F8"

# ---------- helpers ----------
def set_run(r, font=BODY, size=None, bold=None, italic=None, color=None):
    r.font.name = font
    rpr = r._element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:ascii'), font); rf.set(qn('w:hAnsi'), font); rf.set(qn('w:cs'), font)
    if size is not None: r.font.size = Pt(size)
    if bold is not None: r.font.bold = bold
    if italic is not None: r.font.italic = italic
    if color is not None: r.font.color.rgb = color

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def row_flag(row, tag):
    trPr = row._tr.get_or_add_trPr(); trPr.append(OxmlElement('w:' + tag))

def cell_margins(table, t=40, b=40, l=110, r=110):
    tblPr = table._tbl.tblPr
    mar = OxmlElement('w:tblCellMar')
    for tag, w in (('top', t), ('bottom', b), ('left', l), ('right', r)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(w)); e.set(qn('w:type'), 'dxa'); mar.append(e)
    tblPr.append(mar)

def is_toc(t):    return bool(re.search(r"\.{4,}", t)) or "\t" in t
def is_caption(t):return bool(re.match(r"^(Figure|Table)\s+\d", t))

# ---------- 1. styles (fallback for anything not touched directly) ----------
st = d.styles["Normal"]
st.font.name = BODY; st.font.size = Pt(12)
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
st.paragraph_format.line_spacing = 1.3
st.paragraph_format.space_after = Pt(8)
st.paragraph_format.widow_control = True

# ---------- 2. margins (comfortable, slight binding margin) ----------
for sec in d.sections:
    sec.left_margin = Inches(1.0); sec.right_margin = Inches(0.85)
    sec.top_margin = Inches(1.0); sec.bottom_margin = Inches(0.9)

# ---------- 3. find Abstract gate (protect title/front matter above it) ----------
abs_idx = next((i for i, p in enumerate(d.paragraphs) if p.text.strip() == "Abstract"), 0)

def classify(p, i):
    t = p.text.strip(); sn = p.style.name
    if not t: return "empty"
    if i < abs_idx: return "front"
    if is_toc(t): return "toc"
    if is_caption(t): return "caption"
    if re.match(r"^(Chapter|CHAPTER)\s+\d", t): return "chapter"
    if re.match(r"^\d+(\.\d+)+\s+\S", t) and len(t) < 90: return "heading"
    if sn in ("Head", "Heading 1") and len(t) < 90 and not t.endswith((".", ":", ";")): return "heading"
    return "body"

counts = {}
for i, p in enumerate(d.paragraphs):
    k = classify(p, i); counts[k] = counts.get(k, 0) + 1
    pf = p.paragraph_format
    if k == "empty":
        continue
    if k == "front":                         # title page: unify font only, keep layout
        for r in p.runs: r.font.name = HEAD if (p.runs and (r.bold or p.alignment == WD_ALIGN_PARAGRAPH.CENTER)) else BODY
        continue
    if k == "chapter":
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT; pf.keep_with_next = True
        pf.space_before = Pt(6); pf.space_after = Pt(14)
        for r in p.runs: set_run(r, HEAD, 20, True, False, NAVY)
    elif k == "heading":
        dots = (t := p.text.strip()).split()[0].count(".")
        size = {1: 15, 2: 13}.get(dots, 12) if dots else 15
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT; pf.keep_with_next = True
        pf.space_before = Pt(10); pf.space_after = Pt(5)
        for r in p.runs: set_run(r, HEAD, size, True, False, NAVY)
    elif k == "caption":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER; pf.keep_with_next = True
        pf.space_before = Pt(4); pf.space_after = Pt(8)
        for r in p.runs: set_run(r, BODY, 10, True, False, GREY)
    elif k == "toc":
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in p.runs: set_run(r, BODY, 12)
    else:  # body / list
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.3; pf.space_after = Pt(8); pf.widow_control = True
        for r in p.runs: set_run(r, BODY, 12)

# ---------- 4. tables ----------
styled = 0
for tbl in d.tables:
    if tbl.style and "Normal Table" in tbl.style.name:   # skip cover/team layout tables
        continue
    styled += 1
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_margins(tbl)
    for ri, row in enumerate(tbl.rows):
        row_flag(row, "cantSplit")
        if ri == 0:
            row_flag(row, "tblHeader")           # repeat header on each page
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if ri == 0:
                shade(cell, HDR_FILL)
                for p in cell.paragraphs:
                    for r in p.runs: set_run(r, HEAD, 11, True, False, RGBColor(0xFF, 0xFF, 0xFF))
            else:
                if ri % 2 == 0: shade(cell, ZEBRA)
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    for r in p.runs: set_run(r, BODY, 10.5)

d.save(F)
print("classified:", counts)
print("tables styled:", styled)
print("saved:", F)

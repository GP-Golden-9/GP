# -*- coding: utf-8 -*-
"""Stage 1: rebuild TOC / List-of-Figures / List-of-Tables with clean dot-leader
tab stops, split crammed multi-subsection entries into one-per-line, recolor ALL
non-table text black, enable auto-hyphenation. Page numbers filled in stage 2.
"""
import re, shutil, docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

F = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
shutil.copyfile(F, "Swarm_Robot_System_Graduation_Book_CORRECTED.pretoc.docx")
d = docx.Document(F)
BLACK = RGBColor(0, 0, 0)
RIGHT_TAB = Inches(6.42)   # A4 width 8.27 - margins (1.0 + 0.85)
PH = "§"               # placeholder marker for page number (filled in stage 2)

# ---------- 1. auto-hyphenation (tames justified word spacing) ----------
se = d.settings.element
if se.find(qn('w:autoHyphenation')) is None:
    ah = OxmlElement('w:autoHyphenation'); ah.set(qn('w:val'), 'true'); se.append(ah)
    hz = OxmlElement('w:hyphenationZone'); hz.set(qn('w:val'), '360'); se.append(hz)
    nc = OxmlElement('w:doNotHyphenateCaps'); nc.set(qn('w:val'), 'true'); se.append(nc)

# ---------- 2. recolor every non-table run black ----------
for p in d.paragraphs:
    for r in p.runs:
        r.font.color.rgb = BLACK

# ---------- helpers ----------
def insert_after(par):
    new = OxmlElement('w:p'); par._p.addnext(new)
    return Paragraph(new, par._parent)

def parse_label(line):
    line = line.strip()
    if "\t" in line:
        line = line.split("\t")[0]
    line = re.sub(r"\.{2,}.*$", "", line)        # drop dot leaders + trailing number
    line = re.sub(r"\s+\d+\s*$", "", line)        # drop a trailing bare page number
    return line.strip(" .")

def style_entry(p, label, indent_in, bold, size):
    # wipe runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Inches(indent_in)
    pf.line_spacing = 1.12
    pf.space_after = Pt(3); pf.space_before = Pt(6 if bold else 0)
    # reset tab stops -> one right tab with dot leader
    pPr = p._p.get_or_add_pPr()
    for t in pPr.findall(qn('w:tabs')): pPr.remove(t)
    pf.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(label); r1.font.name = "Cambria"; r1.font.size = Pt(size)
    r1.font.bold = bold; r1.font.color.rgb = BLACK
    r2 = p.add_run("\t" + PH); r2.font.name = "Cambria"; r2.font.size = Pt(size)
    r2.font.bold = bold; r2.font.color.rgb = BLACK

def level_of(label):
    m = re.match(r"^(\d+(?:\.\d+)*)", label)
    if not m: return 0
    return label.count(".", 0, m.end())  # number of dots in the leading number

# ---------- 3. rebuild TOC block ----------
def find(txt):
    return next(i for i, p in enumerate(d.paragraphs) if p.text.strip() == txt)

toc_start = find("Table of Contents")
# front-matter entries are the 6 between the heading and 'Chapter 1: Introduction'
fm = ["Acknowledgement", "Abstract", "Table of contents", "List of figures",
      "List of tables", "List of abbreviations and acronyms"]

# Collect the TOC body paragraphs (front-matter + chapter/section lines) as objects.
ps = d.paragraphs
toc_paras = []
for i in range(toc_start + 1, len(ps)):
    t = ps[i].text.strip()
    if t == "List Of Figures":          # end of TOC block
        break
    if t:
        toc_paras.append(ps[i])

# process each TOC paragraph
for p in list(toc_paras):
    raw = p.text
    lines = [l for l in raw.split("\n") if l.strip()]
    if not lines:
        continue
    first = lines[0].strip()
    if first in fm or first.rstrip() in fm:
        style_entry(p, parse_label(lines[0]), 0.0, False, 12); continue
    if re.match(r"^(Chapter|CHAPTER)\s+\d", first):
        style_entry(p, parse_label(lines[0]), 0.0, True, 12.5); continue
    # section / sub-section line(s); first goes into p, extras inserted after
    anchor = p
    for k, ln in enumerate(lines):
        lab = parse_label(ln); lv = level_of(lab)
        indent = 0.25 if lv <= 1 else 0.5
        size = 12 if lv <= 1 else 11
        if k == 0:
            style_entry(p, lab, indent, False, size)
        else:
            np = insert_after(anchor)
            style_entry(np, lab, indent, False, size); anchor = np

# ---------- 4. rebuild List of Figures + List of Tables ----------
def rebuild_list(heading_text):
    idx = find(heading_text)
    j = idx + 1
    while j < len(d.paragraphs):
        t = d.paragraphs[j].text.strip()
        if not t:
            j += 1; continue
        if t.startswith("List ") or t.startswith("Chapter ") or t.startswith("List OF"):
            break
        if re.match(r"^(Figure|Table)\s+\d", t):
            style_entry(d.paragraphs[j], parse_label(t), 0.0, False, 12)
        j += 1

rebuild_list("List Of Figures")
rebuild_list("List of tables")

d.save(F)
print("Stage 1 done: hyphenation on, text black, TOC/LoF/LoT rebuilt with leaders.")
print("TOC paragraphs processed:", len(toc_paras))

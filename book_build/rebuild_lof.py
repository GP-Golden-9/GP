# -*- coding: utf-8 -*-
"""Stage 2: rebuild the List of Figures from the final 37 captions (dot-leader entries),
then recompute page numbers for TOC + List of Figures + List of Tables from the PDF."""
import re, os, json, docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import win32com.client as win32

F = os.path.abspath("Swarm_Robot_System_Graduation_Book_CORRECTED.docx")
PDF = os.path.abspath("Swarm_Robot_System_Graduation_Book_CORRECTED.pdf")
BLACK = RGBColor(0, 0, 0); RIGHT_TAB = Inches(6.42); DASH = "–"; PH = "¤"
ordered = json.load(open("book_build/_figs.json"))["ordered"]

d = docx.Document(F)

def find(txt):
    return next(i for i, p in enumerate(d.paragraphs) if p.text.strip() == txt)

def new_after(par):
    e = OxmlElement('w:p'); par._p.addnext(e); return Paragraph(e, par._parent)

def entry(p, label):
    for r in list(p.runs): r._element.getparent().remove(r._element)
    pf = p.paragraph_format; pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.left_indent = Inches(0.0); pf.line_spacing = 1.12; pf.space_after = Pt(3)
    pPr = p._p.get_or_add_pPr()
    for t in pPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tabs'):
        pPr.remove(t)
    pf.tab_stops.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    r1 = p.add_run(label); r1.font.name = "Cambria"; r1.font.size = Pt(12); r1.font.color.rgb = BLACK
    r2 = p.add_run("\t" + PH); r2.font.name = "Cambria"; r2.font.size = Pt(12); r2.font.color.rgb = BLACK

# --- rebuild List of Figures block: delete old entries, write 37 new ---
lof = find("List Of Figures")
# remove existing entry paragraphs until 'List of tables'
i = lof + 1
while i < len(d.paragraphs):
    t = d.paragraphs[i].text.strip()
    if t == "List of tables":
        break
    if re.match(r"^Figure\s+\d", t) or not t:
        # delete this paragraph
        el = d.paragraphs[i]._p; el.getparent().remove(el)
    else:
        i += 1
# insert fresh entries after the heading
anchor = d.paragraphs[find("List Of Figures")]
for num, title in ordered:
    np = new_after(anchor)
    entry(np, "Figure %s %s %s" % (num, DASH, title))
    anchor = np
d.save(F)
print("List of Figures rebuilt with", len(ordered), "entries")

# --- export, then fill page numbers for every PH placeholder (TOC + LoF + LoT) ---
def export():
    w = win32.DispatchEx("Word.Application"); w.Visible = False; w.DisplayAlerts = 0
    try:
        doc = w.Documents.Open(F, ReadOnly=True)
        doc.ExportAsFixedFormat(OutputFileName=PDF, ExportFormat=17, OpenAfterExport=False,
                                OptimizeFor=0, CreateBookmarks=1)
        doc.Close(False)
    finally:
        w.Quit()

export()
import fitz
pdf = fitz.open(PDF); ptext = [pdf[i].get_text() for i in range(pdf.page_count)]
plow = [t.lower() for t in ptext]

def footer(i):
    ms = re.findall(r"(\d+)\s*\|\s*P", ptext[i]); return int(ms[-1]) if ms else i + 1
def page_of(key, mode="max"):
    hits = [i for i in range(len(plow)) if key.lower() in plow[i]]
    if not hits: return None
    return footer(max(hits) if mode == "max" else min(hits))
def page_caption(kind, num):
    # boundary-aware caption lookup (avoids 'Figure 3.1' matching 3.10-3.17)
    pat = re.compile(r"%s\s+%s\s*[–—:-]" % (kind, re.escape(num)))
    for i in range(len(ptext)):
        if pat.search(ptext[i]): return footer(i)
    return None

FM = {"acknowledgement": ("acknowledgement", "min"), "abstract": ("abstract", "min"),
      "table of contents": ("table of contents", "min"), "list of figures": ("figure 1.1", "min"),
      "list of tables": ("table 1.1", "min"), "list of abbreviations and acronyms": ("list of abbreviation", "min")}

def target(label):
    low = label.lower()
    if low in FM: return page_of(*FM[low])
    if re.match(r"^(chapter|CHAPTER)\s+\d", label, re.I): return page_of(label, "max")
    m = re.match(r"^Figure\s+(\d+\.\d+)", label)
    if m: return page_caption("Figure", m.group(1))
    m = re.match(r"^Table\s+(\d+\.\d+)", label)
    if m: return page_caption("Table", m.group(1))
    if re.match(r"^\d+(\.\d+)*\s", label): return page_of(label, "max")
    return page_of(label, "max")

d = docx.Document(F)
toc0 = find("Table of Contents"); end = find("List OF Abbreviation")
filled = miss = 0; unresolved = []
for p in d.paragraphs[toc0:end]:
    if "\t" not in p.text:        # only dotted-leader entries
        continue
    label = p.text.split("\t")[0].strip()
    if not label:
        continue
    pg = target(label)
    if pg is None:
        miss += 1; unresolved.append(label[:40]); continue
    runs = p.runs
    if runs:
        runs[0].text = label
        if len(runs) > 1:
            runs[1].text = "\t" + str(pg)
            for r in runs[2:]: r.text = ""
        else:
            rr = p.add_run("\t" + str(pg)); rr.font.name = "Cambria"; rr.font.size = Pt(12); rr.font.color.rgb = BLACK
    filled += 1
d.save(F)
print("page numbers recomputed (TOC+LoF+LoT):", filled, "| unresolved:", miss, unresolved[:5])
export()
print("final PDF:", os.path.getsize(PDF), "bytes")

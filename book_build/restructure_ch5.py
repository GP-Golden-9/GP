"""Restructure Chapter 5 results tables (5.1-5.5) so there are NO empty /
"Not measured" cells. Present only data that is genuinely backed:
  - configured design values
  - simulation-soak measurements (docs/baseline/soak_*.json)
  - acceptance-gate targets (tools/soak_test.py)
  - the yaw-drift calibration result (runbook + config/robot2.yaml)
Un-benchmarked quantities are dropped from the tables (their columns removed)
and remain acknowledged in the 5.4.2 limitations text. NO fabricated results.
"""
import shutil
import docx
from docx.oxml.ns import qn

SRC = "Swarm_Robot_System_Graduation_Book_CORRECTED.docx"
BAK = "Swarm_Robot_System_Graduation_Book_CORRECTED.prev.docx"


def delete_column(table, idx):
    grid = table._tbl.tblGrid
    gridcols = grid.findall(qn("w:gridCol"))
    grid.remove(gridcols[idx])
    for row in table.rows:
        tcs = row._tr.findall(qn("w:tc"))
        if idx < len(tcs):
            row._tr.remove(tcs[idx])


def set_row_count(table, n):
    while len(table.rows) < n:
        table.add_row()
    while len(table.rows) > n:
        tr = table.rows[-1]._tr
        table._tbl.remove(tr)


def set_cell(cell, text):
    p0 = cell.paragraphs[0]
    if p0.runs:
        p0.runs[0].text = text
        for r in p0.runs[1:]:
            r.text = ""
    else:
        p0.add_run(text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def fill(table, rows):
    for ri, rowvals in enumerate(rows):
        for ci, val in enumerate(rowvals):
            set_cell(table.rows[ri].cells[ci], val)


def replace_para(doc, find, repl):
    n = 0
    for p in doc.paragraphs:
        full = "".join(r.text for r in p.runs)
        if find in full:
            new = full.replace(find, repl)
            if p.runs:
                p.runs[0].text = new
                for r in p.runs[1:]:
                    r.text = ""
            else:
                p.add_run(new)
            n += 1
    return n


def main():
    shutil.copyfile(SRC, BAK)
    d = docx.Document(SRC)

    # ---- Table 12 = 5.1 SLAM & Mapping (4 cols -> 3) ----
    t = d.tables[12]
    delete_column(t, 3)              # drop "Error Margin"
    set_row_count(t, 5)
    fill(t, [
        ["Mapping Parameter", "Value", "Basis"],
        ["Grid resolution", "0.025 m/cell", "Configured (slam_toolbox)"],
        ["Map publish rate", "<= 1.0 Hz (p95 inter-arrival ~1.0 s in simulation)", "Configured / sim soak"],
        ["SLAM backend", "async slam_toolbox + rf2o laser odometry", "Design"],
        ["Indoor test arena", "4.0 m x 4.0 m", "Physical setup"],
    ])

    # ---- Table 13 = 5.2 Object-detection pipeline (5 cols -> 3) ----
    t = d.tables[13]
    delete_column(t, 4)              # drop "Confidence Threshold" (folded into rows)
    delete_column(t, 3)              # drop "mAP"  -> 5 cols become 3
    set_row_count(t, 7)
    fill(t, [
        ["Detection Pipeline Aspect", "Value", "Basis"],
        ["Primary model", "yolov8s.pt (COCO: person, dog, cat)", "Configured"],
        ["Secondary model", "fire.pt (Fire class only)", "Configured"],
        ["Pipeline throughput", "15.2 fps (simulation soak)", "Measured (sim)"],
        ["Throughput acceptance gate", ">= 12 fps (met)", "tools/soak_test.py"],
        ["Map-marker confidence floors", "person 0.70; dog/cat/fire 0.60", "Configured"],
        ["Audible fire-alarm gate", "0.80", "Configured"],
    ])

    # ---- Table 14 = 5.3 Localization & Odometry calibration (4 cols -> 3) ----
    t = d.tables[14]
    delete_column(t, 3)              # drop "Rotational Drift" col (folded into rows)
    set_row_count(t, 8)
    fill(t, [
        ["Localization Quantity", "Value", "Basis"],
        ["Heading drift, before gyro calibration", "~45 deg over a multi-turn run", "Measured"],
        ["Heading drift, after gyro_scale 1.0304", "~25 deg", "Measured"],
        ["Gyro scale correction", "1.0304 (1048.2 deg read / 1080 deg actual)", "Calibrated"],
        ["Wheel diameter", "0.085 m", "Calibrated"],
        ["Encoder resolution", "408 ticks/rev", "Calibrated"],
        ["Wheel base", "0.225 m (360 deg pivot test)", "Calibrated"],
        ["Heading fusion", "Complementary encoder + gyro, slip-gated", "Design"],
    ])

    # ---- Table 15 = 5.4 Gas detection parameters (5 cols -> 3) ----
    t = d.tables[15]
    delete_column(t, 4)              # drop "Publish Latency"
    delete_column(t, 3)              # drop "Sensor Response Time"  -> 5 cols become 3
    set_row_count(t, 7)
    fill(t, [
        ["Gas Detection Parameter", "Value", "Basis"],
        ["Sensor", "MQ-5 (LPG / natural gas / coal gas)", "Hardware"],
        ["Alarm threshold", "3000 ADC (field-tuned)", "Configured"],
        ["Clear threshold (hysteresis)", "2000 ADC", "Configured"],
        ["Alarm latch", "minimum 10 s", "Configured"],
        ["Poll rate", "3 Hz", "Configured"],
        ["Transport", "HTTP/JSON, polled at the operations center", "Design"],
    ])

    # ---- Table 16 = 5.5 Network metrics (5 cols -> 4) ----
    t = d.tables[16]
    delete_column(t, 2)              # drop "Avg. Bandwidth Usage" (no per-stream tool)
    set_row_count(t, 5)
    fill(t, [
        ["Data Stream", "Protocol / Port", "Measured (simulation soak)", "Acceptance gate"],
        ["Motion control (cmd.drive)", "ZMQ ROUTER/DEALER (5558)", "ACK p95 ~13 ms", "<= 150 ms (met)"],
        ["Sensor telemetry (tele.full)", "ZMQ PUB (5556)", "~19.7 Hz; 0.0% loss (nominal run)", "< 1% seq loss"],
        ["Compressed video", "ZMQ PUB (5560)", "15.2 fps; capture-age within bound", ">= 12 fps (met)"],
        ["SLAM map (map.grid)", "ZMQ PUB (5557)", "inter-arrival p95 ~1.0 s", "<= 2.0 s (met)"],
    ])

    # ---- Captions (body + List of Tables) ----
    caps = [
        ("Table 5.1: SLAM and Mapping Performance Results", "Table 5.1: SLAM and Mapping Configuration"),
        ("Table 5.2: Object Detection Results (primary yolov8s.pt + secondary fire.pt)",
         "Table 5.2: Object-Detection Pipeline Configuration and Throughput"),
        ("Table 5.3: Localization and Odometry Errors", "Table 5.3: Localization and Odometry Calibration"),
        ("Table 5.4: MQ-5 Gas Sensor Performance", "Table 5.4: MQ-5 Gas Detection Parameters"),
    ]
    for old, new in caps:
        print("caption %r -> %d para(s)" % (new, replace_para(d, old, new)))

    # ---- Lead-in sentences that promised measured values we don't report ----
    edits = [
        ("However, translational errors remained visible over longer and more complex paths. These errors were mainly related to wheel slip, floor friction variation, and accumulated odometry uncertainty. Table 5.3 summarizes the localization results obtained during the main trajectory tests.",
         "Some residual drift remained over longer, multi-turn paths, attributable to wheel slip, floor-friction variation, and accumulated odometry uncertainty. Table 5.3 summarizes the localization and odometry calibration that bounds this drift."),
        ("Table 5.4 presents the observed gas sensor response times and ADC thresholds.",
         "Table 5.4 lists the configured MQ-5 detection parameters."),
        ("The publish latency remained low in all hazardous test cases, confirming that the environmental sensing path can support near real-time alerts on the monitoring interface.",
         "With the sensor polled at 3 Hz and alerts raised immediately on threshold crossing, the environmental sensing path supports near real-time alerts on the monitoring interface."),
        ("The isolation of power domains (battery-direct motor drivers and water pump; 5.1V buck regulator for Pi/Mega logic; and 6.0V BEC for Beta's servo) successfully protected the computers from motor-induced voltage brownouts.",
         "The isolation of power domains (battery-direct motor drivers and water pump; a dedicated buck regulator for Pi/Mega logic; and a separate BEC rail for the servo) mitigated motor-induced voltage dips."),
    ]
    for old, new in edits:
        print("prose -> %d para(s): %s..." % (replace_para(d, old, new), new[:50]))

    d.save(SRC)
    print("Saved:", SRC)


if __name__ == "__main__":
    main()
